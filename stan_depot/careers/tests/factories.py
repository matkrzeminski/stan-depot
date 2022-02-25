import pytest
import factory
import factory.fuzzy
from fpdf import FPDF
from docx import Document
from django.template.defaultfilters import slugify


from stan_depot.careers.models import JobOffer
from stan_depot.contact.tests.factories import PlaceFactory


class JobOfferFactory(factory.django.DjangoModelFactory):
    title = factory.fuzzy.FuzzyText()
    slug = factory.LazyAttribute(lambda obj: slugify(obj.title))
    description = factory.Faker("paragraph", nb_sentences=5, variable_nb_sentences=True)
    location = factory.SubFactory(PlaceFactory)
    compensation_min = factory.fuzzy.FuzzyInteger(5000, 10000)
    compensation_max = factory.fuzzy.FuzzyInteger(15000, 25000)
    status = factory.fuzzy.FuzzyChoice(x[0] for x in JobOffer.Status.choices)

    class Meta:
        model = JobOffer


@pytest.fixture
def job_offer():
    return JobOfferFactory()


@pytest.fixture(scope="session")
def generate_pdf(tmp_path_factory):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(40, 10, "Hello World!")
    fn = tmp_path_factory.mktemp("data") / "test.pdf"
    pdf.output(fn)
    return fn


@pytest.fixture(scope="session")
def generate_doc(tmp_path_factory):
    document = Document()
    document.add_heading("This is the title", 0)
    p = document.add_paragraph("And this is text ")
    p.add_run("some bold text").bold = True
    p.add_run("and italic text.").italic = True
    fn = tmp_path_factory.mktemp("data") / "test.doc"
    document.save(fn)
    return fn


@pytest.fixture(scope="session")
def generate_docx(tmp_path_factory):
    document = Document()
    document.add_heading("This is the title", 0)
    p = document.add_paragraph("And this is text ")
    p.add_run("some bold text").bold = True
    p.add_run("and italic text.").italic = True
    fn = tmp_path_factory.mktemp("data") / "test.docx"
    document.save(fn)
    return fn
